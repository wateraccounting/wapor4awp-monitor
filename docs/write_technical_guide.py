"""
write_technical_guide.py

Generates WaPOR4AWP_Global_Technical_Guide.docx using python-docx.

Usage:
    cd global-wapor-awp/docs
    python write_technical_guide.py

Requirements:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from pathlib import Path

OUT = Path(__file__).parent / 'WaPOR4AWP_Global_Technical_Guide.docx'

# ── Colour palette ─────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x00, 0x4B, 0x87)   # IHE Delft blue
BLUE_MID   = RGBColor(0x1A, 0x6F, 0xA8)
TEAL       = RGBColor(0x00, 0x7A, 0x87)
RED_WARN   = RGBColor(0xC0, 0x00, 0x00)
GREY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)
BLACK      = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# ── Page layout ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)


# ── Paragraph helpers ──────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BLUE_DARK
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    doc.add_paragraph()
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE_MID
    p.runs[0].font.size = Pt(13)
    p.runs[0].bold = True
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = TEAL
    p.runs[0].font.size = Pt(11)
    p.runs[0].bold = True
    return p

def para(text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix + '  ')
        run.bold = True
        run.font.color.rgb = BLUE_DARK
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def formula(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(2.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x70)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF3CD')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run('▶ Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x85, 0x60, 0x04)
    p.add_run(text).font.color.rgb = RGBColor(0x85, 0x60, 0x04)
    return p

def warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run('⚠ Warning: ')
    run.bold = True
    run.font.color.rgb = RED_WARN
    p.add_run(text).font.color.rgb = RED_WARN
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._p.get_or_add_pPr().append(shading)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

def table_hdr(tbl, headers, col_widths_cm):
    row = tbl.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shade = OxmlElement('w:shd')
        shade.set(qn('w:val'), 'clear')
        shade.set(qn('w:color'), 'auto')
        shade.set(qn('w:fill'), '004B87')
        cell._tc.get_or_add_tcPr().append(shade)

def table_row(tbl, values, row_idx, col_widths_cm, alt=False):
    row = tbl.add_row()
    fill = 'F2F6FA' if alt else 'FFFFFF'
    for i, (val, w) in enumerate(zip(values, col_widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.paragraphs[0].add_run(str(val)).font.size = Pt(9)
        shade = OxmlElement('w:shd')
        shade.set(qn('w:val'), 'clear')
        shade.set(qn('w:color'), 'auto')
        shade.set(qn('w:fill'), fill)
        cell._tc.get_or_add_tcPr().append(shade)

def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('WaPOR4AWP Global')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = BLUE_DARK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('Annual Agricultural Water Productivity\nfrom WaPOR v3 Satellite Data')
run2.font.size = Pt(16)
run2.font.color.rgb = BLUE_MID

doc.add_paragraph()
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = div.add_run('─' * 50)
run3.font.color.rgb = TEAL

doc.add_paragraph()
type_p = doc.add_paragraph()
type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = type_p.add_run('Technical Guide and Methodological Reference')
run4.font.size = Pt(13)
run4.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

meta_lines = [
    f'Prepared by:  Seleshi G. Yalew',
    f'Affiliation:  IHE Delft Institute for Water Education',
    f'Date:         {datetime.date.today().strftime("%B %Y")}',
    f'Version:      1.0',
]
for line in meta_lines:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.add_run(line).font.size = Pt(11)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')

toc_entries = [
    ('1.', 'Introduction', '3'),
    ('2.', 'System Architecture', '4'),
    ('3.', 'External Data Sources', '5'),
    ('4.', 'Methodology', '7'),
    ('  4.1', 'Core Formula', '7'),
    ('  4.2', 'Effective Precipitation (Pe)', '8'),
    ('  4.3', 'Blue Evapotranspiration (ETb)', '9'),
    ('  4.4', 'GMIE-100 Fractional Weighting', '10'),
    ('  4.5', 'Blue ET Volume (VETb)', '11'),
    ('  4.6', 'Agricultural GVA and Crop Ratio (Cr)', '12'),
    ('  4.7', 'Derived Indicators: cAwp and tAwp', '14'),
    ('5.', 'Processing Pipeline', '15'),
    ('6.', 'Quality Flags', '19'),
    ('7.', 'Dashboard User Guide', '21'),
    ('8.', 'Known Limitations and Uncertainty', '24'),
    ('9.', 'References', '26'),
]

for num, title, pg in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_num = p.add_run(f'{num:<8}')
    run_num.font.bold = num.strip() in [str(i)+'.' for i in range(1,10)]
    run_num.font.size = Pt(10)
    run_title = p.add_run(title)
    run_title.font.bold = num.strip() in [str(i)+'.' for i in range(1,10)]
    run_title.font.size = Pt(10)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1('1.  Introduction')

para(
    'Agricultural Water Productivity (Awp) quantifies how much economic value is generated '
    'per unit of blue water consumed by irrigated agriculture. It is a key metric for '
    'assessing irrigation efficiency, informing water allocation decisions, and tracking '
    'progress toward sustainable water use targets under SDG 6.4.'
)
para(
    'The WaPOR4AWP system was originally developed for Africa and the Middle East and '
    'North Africa (MENA) region using FAO\'s WaPOR v2 satellite data. This document '
    'describes the global extension of that system — WaPOR4AWP Global — which computes '
    'Awp for all countries worldwide using:'
)
bullet('WaPOR v3 monthly evapotranspiration and precipitation data (FAO/IHE Delft)')
bullet('GMIE-100 global irrigation map, used as a fractional pixel weight (Harvard Dataverse)')
bullet('World Bank agricultural GDP (GVA) as the economic numerator')
bullet('AQUASTAT-derived crop ratios to correct GVA for non-irrigated sectors')

doc.add_paragraph()
para(
    'The output is a country-year panel dataset (2018 onwards) of Awp and related '
    'indicators, visualised through a multi-page Streamlit dashboard.'
)

h2('1.1  Scope')
para('Temporal coverage:  2018–present (WaPOR v3 baseline year is 2018)')
para('Spatial coverage:   All countries with GMIE-detected irrigation and World Bank GVA data')
para('Spatial resolution: WaPOR v3 L1 (~300 m) for pixel-level computation; output is country-level')
para('Primary output:     Annual Awp (USD/m³), change indicators, and component variables')

h2('1.2  Relationship to Original WaPOR4AWP')
para(
    'WaPOR4AWP Global differs from the original Africa/MENA system in three important ways:'
)
bullet('Uses WaPOR v3 instead of v2 — scale factors must be re-verified (see Section 5.1)')
bullet('Uses GMIE-100 as a continuous fraction weight (0–1) rather than a binary irrigation mask')
bullet('Runs the entire raster pipeline in Google Earth Engine (GEE) rather than local Python/rasterio, '
       'enabling global coverage without local storage or compute constraints')

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1('2.  System Architecture')

para(
    'The pipeline has three sequential phases. Each phase depends on outputs from the previous one.'
)

h2('2.1  Phase 1 — Google Earth Engine (Raster Processing)')
para('All pixel-level computation runs server-side in GEE, avoiding local data download of global rasters.')

code_block([
    'gee/00_confirm_scale_factors.js   Inspect WaPOR v3 metadata; record SCALE_AETI and SCALE_PCP',
    'gee/01_prepare_gmie_asset.js      Mosaic 67 GMIE-100 tiles; zero background; export GEE asset',
    'gee/02_export_country_awp.js      Compute monthly ETb inline; reduce to country sums;',
    '                                   export one CSV per year to Google Drive',
])

note(
    'Scripts 00 and 01 are run once. Script 02 is re-run each time new WaPOR data becomes '
    'available (e.g., when a new year is released).'
)

h2('2.2  Phase 2 — Python Tabular Processing')
para('Downloads and aligns socioeconomic data, then merges with GEE outputs to compute Awp.')

code_block([
    'scripts/02_download_worldbank_gva.py   Download World Bank NV.AGR.TOTL.CD via API',
    'scripts/04_build_country_crosswalk.py  Build ISO-3 master alignment table',
    'scripts/03_prepare_aquastat_cr.py      Derive Cr from AQUASTAT; apply fallback hierarchy',
    'scripts/05_compute_annual_awp.py       Merge GEE + WB + Cr; compute Awp, cAwp, tAwp; flags',
    'scripts/06_validate_outputs.py         Automated QA; generate validation_report.txt',
])

note('Scripts 02 and 04 are independent of each other and can run in parallel. '
     'Script 03 requires the crosswalk from 04. Script 05 requires all four.')

h2('2.3  Phase 3 — Streamlit Dashboard')
para('Reads the final CSV and GeoJSON. No database or server required.')

code_block([
    'dashboard/streamlit_app.py                Entry point',
    'dashboard/utils.py                        Shared constants, indicators, helpers',
    'dashboard/pages/1_Global_Overview.py      Choropleth world map',
    'dashboard/pages/2_Country_Explorer.py     Time-series per country',
    'dashboard/pages/3_Rankings_and_Comparison.py  Country and regional rankings',
    'dashboard/pages/4_Methodology_and_Download.py  Documentation and CSV export',
])

h2('2.4  Data Flow Diagram')
code_block([
    'WaPOR v3 (GEE)  ──┐',
    'GMIE-100 (GEE)  ──┼──► [GEE: 02_export_country_awp.js] ──► gee_exports/*.csv',
    '                ──┘                                               │',
    'World Bank API  ──────► [02_download_worldbank_gva.py]           │',
    'AQUASTAT (manual) ────► [03_prepare_aquastat_cr.py]  ────────────┤',
    'WB API (country)  ────► [04_build_country_crosswalk.py]          │',
    '                                                                  ▼',
    '                                              [05_compute_annual_awp.py]',
    '                                                       │',
    '                                          data/final/global_awp_country_year.csv',
    '                                                       │',
    '                                       [Streamlit dashboard]',
])

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. EXTERNAL DATA SOURCES
# ══════════════════════════════════════════════════════════════════════════════
h1('3.  External Data Sources')

h2('3.1  WaPOR v3 — Evapotranspiration and Precipitation')

para(
    'WaPOR (Water Productivity through Open access of Remotely sensed derived data) is a '
    'FAO initiative developed in collaboration with IHE Delft. Version 3 provides global '
    'coverage at Level 1 (~300 m spatial resolution).'
)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
table_hdr(tbl, ['Product', 'GEE Collection ID', 'Temporal', 'Unit'], [3.5, 7.0, 2.5, 2.5])
rows = [
    ('L1-AETI-M', 'projects/UNFAO/wapor/v3/L1-AETI-M', 'Monthly', 'mm/month (raw × scale factor)'),
    ('L1-PCP-M',  'projects/UNFAO/wapor/v3/L1-PCP-M',  'Monthly', 'mm/month (raw × scale factor)'),
]
for i, r in enumerate(rows):
    table_row(tbl, r, i, [3.5, 7.0, 2.5, 2.5], alt=bool(i % 2))

doc.add_paragraph()
warning(
    'WaPOR v3 scale factors must be verified empirically before any computation by '
    'running gee/00_confirm_scale_factors.js. Do not assume the WaPOR v2 scale factor '
    '(×0.1) applies to v3.'
)

h2('3.2  GMIE-100 — Global Map of Irrigation Areas')
para(
    'GMIE-100 (Global Map of Irrigation Extent at 100 arcseconds ≈ 1 km) provides a '
    'fractional measure of irrigated area per grid cell. Values range from 0 (no '
    'irrigation) to 1 (fully irrigated), representing the proportion of each pixel '
    'equipped for or under irrigation.'
)
para('Source:        Harvard Dataverse, doi:10.7910/DVN/HKBAQQ')
para('Format:        67 GeoTIFF tiles; background value = −99')
para('Resolution:    ~1 km (100 arcseconds)')
para('Vintage:       Circa 2000; static — does not change year to year')
para('Access:        Manual download from Dataverse; tiles uploaded to GEE by user')

note(
    'GMIE-100 is used as a continuous fraction (weight), not a binary mask. A pixel '
    'with GMIE = 0.4 contributes 40% of its water volume and area to the country total. '
    'Background pixels (−99) are set to 0, not masked, so that weighted sums over '
    'country boundaries remain numerically stable.'
)

h2('3.3  World Bank — Agricultural GVA')
para(
    'The economic numerator for Awp is sourced from the World Bank World Development '
    'Indicators (WDI) via the public API.'
)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
table_hdr(tbl2, ['Indicator code', 'Description', 'Access'], [3.5, 8.0, 4.0])
rows2 = [
    ('NV.AGR.TOTL.CD', 'Agriculture, forestry and fishing value added (current US$)',
     'api.worldbank.org/v2/country/all/indicator/...'),
    ('/v2/country', 'Country metadata: name, ISO-3, World Bank region, income group',
     'api.worldbank.org/v2/country'),
]
for i, r in enumerate(rows2):
    table_row(tbl2, r, i, [3.5, 8.0, 4.0], alt=bool(i % 2))

doc.add_paragraph()
warning(
    'NV.AGR.TOTL.CD includes agriculture, forestry, AND fishing. For countries '
    'with large non-crop sectors (e.g., Norway, Iceland, Indonesia), this inflates '
    'the Awp numerator. Such countries are flagged GVA_FORESTRY_FISHING.'
)

h2('3.4  AQUASTAT — Crop Ratio (Cr) Data')
para(
    'The crop ratio Cr is the fraction of World Bank agricultural GVA that is '
    'attributable to non-irrigated production. It corrects the GVA numerator to '
    'represent irrigated crop value only. AQUASTAT data are used to derive Cr from '
    'irrigated and cultivated area statistics.'
)
para('Source:   FAO AQUASTAT Statistics Portal (fao.org/aquastat/statistics/query)')
para('Access:   Manual download — requires selection of variables and country list')
para(
    'Variables needed: "Total area equipped for irrigation" (ha) and '
    '"Cultivated area" (ha), all countries, 2010–present'
)

note(
    'AQUASTAT data are optional. If the expected CSV files are absent from '
    'data/raw/aquastat/, the pipeline uses country-fixed Cr values from literature '
    '(embedded in scripts/03_prepare_aquastat_cr.py) with regional and global defaults '
    'as further fallbacks. See Section 4.6 for the full fallback hierarchy.'
)

h2('3.5  Country Boundaries')

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = 'Table Grid'
table_hdr(tbl3, ['Dataset', 'Where used', 'Access', 'Property for ISO-3'], [4, 4, 4, 3.5])
rows3 = [
    ('USDOS LSIB_SIMPLE 2017', 'GEE reduceRegions (country aggregation)',
     'GEE built-in: USDOS/LSIB_SIMPLE/2017', 'iso_alpha3'),
    ('global_boundaries_iso3.geojson', 'Streamlit choropleth map',
     'Pre-existing project file', 'iso3'),
]
for i, r in enumerate(rows3):
    table_row(tbl3, r, i, [4, 4, 4, 3.5], alt=bool(i % 2))

doc.add_paragraph()
note(
    'These are two separate boundary datasets. ISO-3 code alignment between them '
    'is handled by the country crosswalk (scripts/04_build_country_crosswalk.py). '
    'Known mismatches are flagged as BOUNDARY_MISMATCH.'
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
h1('4.  Methodology')

h2('4.1  Core Formula')

para(
    'Agricultural Water Productivity is defined as the value of irrigated agricultural '
    'production per unit volume of blue water consumed:'
)
formula('Awp = GVAa × (1 − Cr) / VETb          [USD/m³]')

doc.add_paragraph()
tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
table_hdr(tbl4, ['Symbol', 'Description', 'Unit'], [2.5, 10.0, 3.0])
rows4 = [
    ('Awp',  'Agricultural Water Productivity',                              'USD/m³'),
    ('GVAa', 'Agriculture, forestry and fishing value added (World Bank)',   'USD/year'),
    ('Cr',   'Crop ratio — non-irrigated fraction of GVAa',                  'dimensionless [0, 1]'),
    ('VETb', 'Annual blue evapotranspiration volume over irrigated land',    'm³/year'),
]
for i, r in enumerate(rows4):
    table_row(tbl4, r, i, [2.5, 10.0, 3.0], alt=bool(i % 2))

doc.add_paragraph()
para(
    'The term GVAa × (1 − Cr) approximates the economic value attributable specifically '
    'to irrigated crop production, by removing the contribution of non-irrigated farming, '
    'forestry, and fishing from the World Bank aggregate.'
)

h2('4.2  Effective Precipitation (Pe)')

para(
    'Effective precipitation (Pe) is the portion of monthly rainfall that contributes '
    'to crop water requirements — i.e., the portion that reduces the need for irrigation. '
    'It is computed using the Brouwer and Heibloem (1986) piecewise formula:'
)

formula('If P ≤ 16.67 mm/month:        Pe = 0')
formula('If 16.67 < P ≤ 75 mm/month:  Pe = 0.6 × P − 10')
formula('If P > 75 mm/month:           Pe = 0.8 × P − 25')

doc.add_paragraph()
para(
    'where P is monthly precipitation in mm/month and Pe is effective precipitation '
    'in mm/month. The formula is piecewise linear with two properties that make it '
    'physically consistent:'
)
bullet(
    'Continuity at breakpoints: at P = 16.67 mm, both segments give Pe = 0. '
    'At P = 75 mm, both segments give Pe = 35 mm. There are no discontinuities.'
)
bullet(
    'Zero threshold: below P = 16.67 mm/month, rainfall is too low to contribute '
    'meaningfully to soil water storage — Pe = 0. This represents months with very '
    'low precipitation (arid environments or dry seasons).'
)
bullet(
    'Diminishing returns at high rainfall: the slope decreases from 0.6 to 0.8 at '
    'P > 75 mm (i.e., more precipitation runs off rather than being stored).'
)

note(
    'Pe is applied per pixel per month before summing to annual ETb. This is more '
    'accurate than applying Pe to annual precipitation totals, because the non-linear '
    'formula responds differently at different monthly timescales.'
)

h2('4.3  Blue Evapotranspiration (ETb)')

para(
    'Blue evapotranspiration (ETb) is the component of actual evapotranspiration (AETI) '
    'that originates from irrigation water (blue water) rather than direct rainfall. '
    'It is computed per pixel per month as:'
)
formula('ETb_month = max(AETI_month − Pe_month, 0)          [mm/month]')

doc.add_paragraph()
para('The annual depth is the sum across all available months:')
formula('ETb_annual = Σ(m=1 to 12) ETb_month               [mm/year]')

doc.add_paragraph()
para(
    'This approach assumes that when AETI exceeds Pe, the deficit is met by irrigation. '
    'When AETI ≤ Pe (rainfed conditions dominate), ETb = 0 — there is no blue water '
    'consumption at that pixel in that month.'
)
para(
    'Both AETI and PCP are sourced from WaPOR v3 products L1-AETI-M and L1-PCP-M '
    'respectively, at monthly temporal resolution and ~300 m spatial resolution. '
    'PCP in WaPOR v3 is derived from CHIRPS but accessed via the WaPOR API/GEE '
    'collection to maintain consistency with the AETI product.'
)

warning(
    'If WaPOR PCP data is missing for a given month (rare but possible in near-real-time '
    'updates), Pe is assumed to be zero for that month, meaning ETb = AETI. This is a '
    'conservative assumption (overestimates blue ET). Missing PCP months are tracked in '
    'the n_months_pcp field and trigger the PARTIAL_YEAR quality flag.'
)

h2('4.4  GMIE-100 Fractional Weighting')

para(
    'GMIE-100 provides, for each ~1 km grid cell, the fraction of the cell that is '
    'equipped for or under irrigation (fGMIE ∈ [0, 1]). This is used as a continuous '
    'weight rather than a binary mask, so that partially irrigated pixels contribute '
    'proportionally to country totals.'
)

para('Before use in computation, GMIE undergoes three preprocessing steps:')
bullet('Background removal: cells with value = −99 (non-land background) are set to 0')
bullet(
    'Bilinear resampling from ~1 km to WaPOR 300 m CRS, so that the GMIE weight '
    'grid aligns with the AETI and PCP pixel grid'
)
bullet(
    'Minimum threshold masking: pixels with fGMIE < 0.01 are excluded from computation. '
    'This removes effectively non-irrigated pixels and reduces global pixel count by '
    '~70–80%, the most important scalability optimisation'
)

formula('GMIE_contribution = fGMIE × (pixel quantity)')

doc.add_paragraph()
note(
    'Using GMIE as a continuous weight rather than a binary mask is critical for '
    'accuracy. A pixel at the boundary of an irrigated perimeter may have fGMIE = 0.3 — '
    'meaning 30% of the pixel is irrigated. Treating it as 0 (excluded) or 1 (fully '
    'irrigated) would both introduce systematic errors in VETb.'
)

h2('4.5  Blue ET Volume (VETb)')

para(
    'VETb is the annual volume of blue evapotranspiration consumed by irrigation over '
    'a country\'s irrigated land, computed from pixel-level ETb and GMIE:'
)

formula('VETb_pixel = (ETb_annual / 1000) × Apixel × fGMIE   [m³/pixel/year]')
formula('VETb_country = Σ VETb_pixel over all pixels in country  [m³/year]')

doc.add_paragraph()
para(
    'where ETb_annual / 1000 converts depth from mm to metres, and Apixel is the '
    'pixel area in m². Pixel area is computed with GEE\'s ee.Image.pixelArea(), '
    'which accounts for the latitude-dependent area of geographic (lon/lat) pixels — '
    'a pixel at 60°N is roughly half the area of a pixel at 0°.'
)

para('Two area-weighted mean depth fields are also computed:')
formula('AETI_depth = Σ(AETI × Apixel × fGMIE) / Σ(Apixel × fGMIE)   [mm/year]')
formula('ETb_depth  = Σ(ETb  × Apixel × fGMIE) / Σ(Apixel × fGMIE)   [mm/year]')

doc.add_paragraph()
para(
    'The GMIE-weighted irrigated area is also exported:'
)
formula('GMIE_area = Σ(Apixel × fGMIE) / 10000                         [ha]')

doc.add_paragraph()
note(
    'VETb is stored in the output CSV as m³ for full precision. The dashboard '
    'displays it as Mm³ (÷ 1,000,000) for readability. The column vetb_Mm3 is '
    'derived automatically from VETb_m3 and is never stored independently.'
)

h2('4.6  Agricultural GVA and Crop Ratio (Cr)')

h3('4.6.1  GVA Source')
para(
    'The economic numerator GVAa is the World Bank indicator NV.AGR.TOTL.CD — '
    '"Agriculture, forestry and fishing, value added, current US$". It is the '
    'broadest agricultural GDP measure and includes crop farming, livestock, '
    'forestry, and fishing in a single figure.'
)

h3('4.6.2  Crop Ratio Definition')
para(
    'Because GVAa includes forestry and fishing (which do not consume irrigation '
    'water), a crop ratio Cr is applied to isolate the irrigated-crop contribution:'
)
formula('Corrected numerator = GVAa × (1 − Cr)')
formula('Cr = fraction of GVAa that is NOT from irrigated crops  ∈ [0, 1]')

doc.add_paragraph()
para(
    'In practice, Cr is estimated from AQUASTAT irrigated and cultivated area statistics '
    'as a proxy:'
)
formula('Cr ≈ 1 − (irrigated_area / cultivated_area)')

doc.add_paragraph()
warning(
    'This is a proxy, not a direct economic measurement. The true Cr would require '
    'country-level data on irrigated crop value, forestry value, and fishing value '
    'separately — data that are not consistently available globally. Cr uncertainty '
    'is typically ±15–25% of the computed Awp value.'
)

h3('4.6.3  Fallback Hierarchy')
para('Cr is assigned using the following priority order:')

tbl5 = doc.add_table(rows=1, cols=4)
tbl5.style = 'Table Grid'
table_hdr(tbl5, ['Level', 'Source', 'Quality flag', 'Coverage'], [1.5, 6.0, 3.5, 4.5])
rows5 = [
    ('1', 'AQUASTAT annual irrigated/cultivated area ratio',
     '(none — highest confidence)', 'Countries with AQUASTAT time series'),
    ('2', 'Country-fixed Cr from literature (embedded in code)',
     'CR_COUNTRY_FIXED', '~35 key countries'),
    ('3', 'World Bank regional average Cr',
     'CR_REGIONAL_DEFAULT', 'All countries with WB region'),
    ('4', 'Global default Cr = 0.80',
     'CR_GLOBAL_DEFAULT', 'Any remaining countries'),
]
for i, r in enumerate(rows5):
    table_row(tbl5, r, i, [1.5, 6.0, 3.5, 4.5], alt=bool(i % 2))

doc.add_paragraph()

h3('4.6.4  Regional Default Cr Values')
tbl6 = doc.add_table(rows=1, cols=2)
tbl6.style = 'Table Grid'
table_hdr(tbl6, ['World Bank Region', 'Default Cr'], [10.5, 5.0])
rows6 = [
    ('Sub-Saharan Africa',         '0.85  (mostly rainfed smallholders)'),
    ('Middle East & North Africa', '0.55  (irrigation-intensive region)'),
    ('South Asia',                 '0.62  (large irrigated sector)'),
    ('East Asia & Pacific',        '0.72'),
    ('Europe & Central Asia',      '0.78'),
    ('Latin America & Caribbean',  '0.80  (mostly rainfed)'),
    ('North America',              '0.76'),
]
for i, r in enumerate(rows6):
    table_row(tbl6, r, i, [10.5, 5.0], alt=bool(i % 2))

doc.add_paragraph()

h2('4.7  Derived Indicators: cAwp and tAwp')

para(
    'Two change indicators are computed from the annual Awp time series to express '
    'short-term dynamics and long-term trends:'
)

h3('4.7.1  Year-to-Year Change (cAwp)')
formula('cAwp_t = (Awp_t − Awp_{t−1}) / Awp_{t−1} × 100          [%]')
para(
    'cAwp measures the percentage change in Awp from the previous year. It is undefined '
    '(NaN) for the baseline year 2018 because no prior-year Awp exists. A positive cAwp '
    'indicates improving water productivity; negative indicates decline.'
)

h3('4.7.2  Trend Since Baseline (tAwp)')
formula('tAwp_t = (Awp_t − Awp_2018) / Awp_2018 × 100             [%]')
para(
    'tAwp measures the percentage change relative to the 2018 baseline. By definition, '
    'tAwp = 0 for the year 2018. It is undefined for countries where Awp could not be '
    'computed in 2018 (e.g., missing GVA or insufficient GMIE area).'
)

note(
    'The baseline year is 2018 because it is the first year for which WaPOR v3 '
    'provides a complete 12-month AETI and PCP record. If the WaPOR v3 archive '
    'is extended backwards, the baseline year may be updated — all tAwp values '
    'must then be recomputed.'
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. PROCESSING PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
h1('5.  Processing Pipeline')

para(
    'This section describes each script in execution order, its inputs, outputs, '
    'key parameters, and what to check after running it.'
)

h2('5.1  GEE Script 00: Confirm Scale Factors')
para('File: gee/00_confirm_scale_factors.js')
para('Run: once, before any other computation')
para(
    'Purpose: WaPOR v3 stores raw digital numbers that must be multiplied by a scale '
    'factor to obtain physical units (mm/month). The scale factor may differ from '
    'WaPOR v2. This script prints all image metadata to the GEE Console so you can '
    'read off the correct factor.'
)
para('Outputs printed to GEE Console (not exported):')
bullet('Band names for L1-AETI-M and L1-PCP-M')
bullet('Scale factor and offset (from image properties)')
bullet('Nominal pixel scale in metres')
bullet('Sample pixel values at Nile Delta (31.2°E, 30.5°N) — cross-check against known values')
bullet('Min/max/mean over 50 km box at Nile Delta')

para('Action required:')
bullet('Record SCALE_AETI and SCALE_PCP from console output')
bullet('Expected July AETI at Nile Delta: ~80–150 mm/month after scaling')
bullet('Update the SCALE_AETI and SCALE_PCP constants at the top of gee/02_export_country_awp.js')

warning('Do not skip this step. Running the pipeline with wrong scale factors produces '
        'ETb values off by 10× or 100×, which would not be caught until the Awp sanity check.')

h2('5.2  GEE Script 01: Prepare GMIE Asset')
para('File: gee/01_prepare_gmie_asset.js')
para('Run: once (GMIE-100 is static)')
para(
    'Purpose: the 67 GMIE-100 GeoTIFF tiles must first be uploaded to GEE as individual '
    'image assets, then this script mosaics them, zeroes out background pixels (−99), '
    'and exports a single global asset.'
)

para('Prerequisites:')
code_block([
    '# Upload all 67 tiles (example for one tile):',
    'earthengine upload image \\',
    '  --asset_id=projects/YOUR_PROJECT/assets/gmie100_tiles/tile_101W_30N \\',
    '  GMIE-100_101W_30N.tif',
    '',
    '# Repeat for each tile. Tiles are in data/raw/gmie/',
])

para('What the script does:')
bullet('Loads the tile collection from projects/YOUR_PROJECT/assets/gmie100_tiles/')
bullet('Mosaics with ee.ImageCollection.mosaic() — first non-null value per pixel')
bullet('Replaces all values < 0 with 0 (zeroes out the −99 background)')
bullet('Clamps to [0, 1] as a safety measure')
bullet('Exports to projects/YOUR_PROJECT/assets/gmie100_global at 1000 m resolution')

para('Sanity checks printed to Console:')
bullet('Mean GMIE at Nile Delta 50 km box: expect 0.7–1.0')
bullet('Mean GMIE in Sahara 50 km box: expect ~0')

h2('5.3  GEE Script 02: Main Pipeline — Export Country Summaries')
para('File: gee/02_export_country_awp.js')
para('Run: once per year range; re-run each year as new WaPOR data is released')

para(
    'This is the core raster computation script. It performs the full pixel-level '
    'ETb computation inline (no intermediate asset storage), then reduces to country '
    'boundaries and exports one CSV per year to Google Drive.'
)

para('Key parameters in the configuration block:')
tbl7 = doc.add_table(rows=1, cols=3)
tbl7.style = 'Table Grid'
table_hdr(tbl7, ['Parameter', 'Default', 'Notes'], [3.5, 2.5, 9.5])
rows7 = [
    ('PROJECT_ID',        '(required)',  'Your GEE cloud project ID'),
    ('SCALE_AETI',        '0.1',        'Update from script 00 output'),
    ('SCALE_PCP',         '0.1',        'Update from script 00 output'),
    ('YEAR_START',        '2018',       'First year of WaPOR v3'),
    ('YEAR_END',          '2023',       'Update as new WaPOR years are released'),
    ('GMIE_MIN_FRACTION', '0.01',       'Pixels below excluded; reduces compute by ~75%'),
    ('TILE_SCALE',        '4',          'Increase to 8 or 16 if tasks time out or OOM'),
]
for i, r in enumerate(rows7):
    table_row(tbl7, r, i, [3.5, 2.5, 9.5], alt=bool(i % 2))

doc.add_paragraph()
para('Output columns per exported CSV:')

tbl8 = doc.add_table(rows=1, cols=3)
tbl8.style = 'Table Grid'
table_hdr(tbl8, ['Column', 'Type', 'Description'], [3.5, 2.0, 10.0])
rows8 = [
    ('iso3',           'string', 'ISO 3166-1 alpha-3 country code'),
    ('country_name',   'string', 'Country name from LSIB boundary'),
    ('year',           'integer','Calendar year'),
    ('n_months_aeti',  'integer','Number of AETI months available (12 = complete year)'),
    ('n_months_pcp',   'integer','Number of PCP months available'),
    ('VETb_m3',        'float', 'GMIE-weighted blue ET volume sum (m³/year)'),
    ('gmie_area_m2',   'float', 'GMIE-weighted area sum (m²); divide by 10000 for ha'),
    ('AETI_wsum',      'float', 'Weighted AETI numerator (mm × m² × GMIE); use for depth mean'),
    ('ETb_wsum',       'float', 'Weighted ETb numerator (mm × m² × GMIE); use for depth mean'),
]
for i, r in enumerate(rows8):
    table_row(tbl8, r, i, [3.5, 2.0, 10.0], alt=bool(i % 2))

doc.add_paragraph()
note(
    'Multi-polygon countries (e.g. Indonesia, Philippines, USA with islands) produce '
    'multiple rows per year in the GEE export — one per polygon feature in LSIB_SIMPLE. '
    'Python script 05 aggregates these by iso3 before computing Awp.'
)

para('If a GEE task fails:')
tbl9 = doc.add_table(rows=1, cols=2)
tbl9.style = 'Table Grid'
table_hdr(tbl9, ['Error message', 'Action'], [6.5, 9.0])
rows9 = [
    ('"Computation timed out"',         'Increase TILE_SCALE to 8 in the config block'),
    ('"User memory limit exceeded"',    'Increase TILE_SCALE to 16'),
    ('"Too many pixels"',               'Already handled by GMIE_MIN_FRACTION mask'),
    ('"Asset not found" for GMIE',      'Verify GMIE_ASSET path; ensure script 01 completed'),
    ('"Image.load: Asset not found"',   'Check WaPOR collection ID; verify GEE project access'),
]
for i, r in enumerate(rows9):
    table_row(tbl9, r, i, [6.5, 9.0], alt=bool(i % 2))

doc.add_paragraph()

h2('5.4  Python Script: Download World Bank GVA')
para('File: scripts/02_download_worldbank_gva.py')
para('Run: once; re-run to refresh for new years')
para(
    'Downloads NV.AGR.TOTL.CD for all countries, 2010–2024, paginating through '
    'the World Bank API. Filters out aggregate regions (codes with length ≠ 3). '
    'Outputs worldbank_gva_agriculture.csv and a coverage report.'
)
para('Key sanity check: prints Ethiopia GVA time series — verify values are in '
     'the expected range (billions of USD) and the trend is consistent.')

h2('5.5  Python Script: Build Country Crosswalk')
para('File: scripts/04_build_country_crosswalk.py')
para('Run: once; re-run if boundary file or WB data changes')
para(
    'Aligns ISO-3 codes across World Bank API, the dashboard GeoJSON boundary file, '
    'and known AQUASTAT conventions. Flags 12 known problem cases (Kosovo, Taiwan, '
    'Western Sahara, South Sudan, etc.) for manual review.'
)
para('Outputs: country_crosswalk.csv with include_flag ∈ {YES, REVIEW, NO}')

h2('5.6  Python Script: Prepare Crop Ratio (Cr)')
para('File: scripts/03_prepare_aquastat_cr.py')
para('Requires: country_crosswalk.csv from script 04')
para(
    'Builds a complete country-year Cr table applying the four-level fallback '
    'hierarchy described in Section 4.6.3. If AQUASTAT files are present in '
    'data/raw/aquastat/, they are used at level 1. Otherwise, the script proceeds '
    'directly to country-fixed and regional defaults.'
)

h2('5.7  Python Script: Compute Annual Awp')
para('File: scripts/05_compute_annual_awp.py')
para('Requires: GEE exports, World Bank GVA, Cr table, crosswalk')
para(
    'This is the main tabular integration step. It:'
)
bullet('Loads and aggregates GEE CSV exports (sums multi-polygon rows by iso3×year)')
bullet('Computes AETI_annual_mm and ETb_annual_mm as area-weighted means from wsum columns')
bullet('Merges World Bank GVA and Cr by iso3×year')
bullet('Computes Awp, cAwp, tAwp')
bullet('Assigns quality flags')
bullet('Writes data/final/global_awp_country_year.csv')

h2('5.8  Python Script: Validate Outputs')
para('File: scripts/06_validate_outputs.py')
para('Run: after every pipeline execution')
para('Runs 10 automated checks and writes data/final/validation_report.txt. '
     'Returns exit code 0 on full pass, 1 if any check fails.')
para('Key checks: ISO-3 coverage, year completeness, Awp range, VETb unit, '
     'Cr range, tAwp = 0 at 2018, no duplicate rows, known-country plausibility.')

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. QUALITY FLAGS
# ══════════════════════════════════════════════════════════════════════════════
h1('6.  Quality Flags')

para(
    'Every row in global_awp_country_year.csv carries a quality_flag field containing '
    'one or more comma-separated flags. Multiple flags can apply simultaneously. '
    'Awp is set to NaN (not computed) when any of the critical data inputs are absent '
    'or zero; quality flags explain why.'
)

tbl10 = doc.add_table(rows=1, cols=3)
tbl10.style = 'Table Grid'
table_hdr(tbl10, ['Flag', 'Awp computed?', 'Description'], [4.5, 2.0, 9.0])
rows10 = [
    ('OK',                    'Yes',   'All key inputs available; indicator calculated normally'),
    ('MISSING_GVA',           'No',    'World Bank GVA missing for this country-year'),
    ('MISSING_CR',            'No',    'No Cr available at any fallback level (very rare)'),
    ('CR_COUNTRY_FIXED',      'Yes',   'Country fixed Cr from literature used (no annual value)'),
    ('CR_REGIONAL_DEFAULT',   'Yes',   'Regional average Cr used — moderate confidence'),
    ('CR_GLOBAL_DEFAULT',     'Yes',   'Global default Cr = 0.80 used — low confidence'),
    ('ZERO_VETB',             'No',    'VETb ≤ 0 — no computable irrigation water use'),
    ('LOW_GMIE_AREA',         'No',    'GMIE-weighted area below 1,000 ha threshold'),
    ('NO_GMIE_IRRIGATION',    'No',    'No GMIE irrigation fraction detected in this country'),
    ('GVA_FORESTRY_FISHING',  'Yes',   'Large forestry/fishing sector — GVA likely overestimates irrigated crop value'),
    ('BOUNDARY_MISMATCH',     'Yes*',  'ISO-3 mismatch between boundary and GEE datasets suspected'),
    ('PARTIAL_YEAR',          'Yes*',  'n_months_aeti < 12 — annual ETb may be underestimated'),
    ('OUTLIER_AWP',           'Yes',   'Awp > 50 or < 0.001 USD/m³ — review recommended'),
]
for i, r in enumerate(rows10):
    table_row(tbl10, r, i, [4.5, 2.0, 9.0], alt=bool(i % 2))

doc.add_paragraph()
para('* Awp is computed but results should be interpreted with caution.')

h2('6.1  Recommended Data Filters by Use Case')

tbl11 = doc.add_table(rows=1, cols=2)
tbl11.style = 'Table Grid'
table_hdr(tbl11, ['Use case', 'Recommended quality_flag filter'], [6.0, 9.5])
rows11 = [
    ('Peer-reviewed scientific publication', 'OK only'),
    ('Country-level policy planning',         'OK + CR_COUNTRY_FIXED'),
    ('Regional or global overview map',       'OK + CR_COUNTRY_FIXED + CR_REGIONAL_DEFAULT'),
    ('Exploratory analysis (flag caveats)',   'All flags (with disclosures)'),
    ('Strictly exclude incomplete years',     'Exclude PARTIAL_YEAR regardless of other flags'),
]
for i, r in enumerate(rows11):
    table_row(tbl11, r, i, [6.0, 9.5], alt=bool(i % 2))

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. DASHBOARD USER GUIDE
# ══════════════════════════════════════════════════════════════════════════════
h1('7.  Dashboard User Guide')

h2('7.1  Running the Dashboard')
code_block([
    'cd global-wapor-awp',
    'pip install -r requirements.txt',
    'streamlit run dashboard/streamlit_app.py',
])
para('The dashboard opens at http://localhost:8501 in your default browser.')
note(
    'The dashboard reads data/final/global_awp_country_year.csv and '
    'data/raw/boundaries/global_boundaries_iso3.geojson from the project directory. '
    'If either file is missing, an error message is shown with instructions to run '
    'the processing pipeline first.'
)

h2('7.2  Page 1 — Global Overview')
para(
    'Displays a choropleth world map for a selected indicator and year.'
)

tbl12 = doc.add_table(rows=1, cols=2)
tbl12.style = 'Table Grid'
table_hdr(tbl12, ['Control', 'Description'], [4.5, 11.0])
rows12 = [
    ('Year slider',          'Select any year from the available range (2018–present)'),
    ('Indicator dropdown',   'Choose from 8 indicators: Awp, cAwp, tAwp, VETb, GVA, GMIE area, ETb depth, AETI depth'),
    ('OK quality only',      'When checked, only rows with quality_flag = OK are shown on the map'),
]
for i, r in enumerate(rows12):
    table_row(tbl12, r, i, [4.5, 11.0], alt=bool(i % 2))

doc.add_paragraph()
para('Metrics row above the map shows: selected year, countries with data, '
     'median Awp (or total VETb/GVA), and highest-value country.')

h2('7.3  Page 2 — Country Explorer')
para(
    'Compares time-series data for up to ~10 selected countries.'
)
bullet('Primary chart: line plot of selected indicator over 2018–present')
bullet('cAwp and tAwp: side-by-side line charts with zero reference line')
bullet('"Show components" checkbox: adds VETb and GVA bar charts for selected countries')
bullet('Data table: full record for selected countries, downloadable via browser')

h2('7.4  Page 3 — Rankings and Comparison')
para('Country rankings for a selected year and indicator.')
bullet('Top N / Bottom N bar charts (N configurable 5–50), coloured by World Bank region')
bullet('Regional box plots: distribution of the indicator across all countries in each region')
bullet('Income group box plots: distribution by World Bank income classification')
bullet('Scatter plot (Awp page only): Awp vs VETb with bubble size = GMIE area')
bullet('Filters: region, income group, quality flag — independently stackable')

h2('7.5  Page 4 — Methodology and Download')
para('Three tabs:')
bullet(
    'Methodology: full formula derivation in rendered LaTeX, data source table, '
    'known limitations summary'
)
bullet(
    'Quality flags: complete flag table with descriptions, filtering guidance'
)
bullet(
    'Download: configurable CSV export — filter by year, quality flag, and column set. '
    'Provides a Download CSV button; no server upload required'
)

h2('7.6  Available Indicators')

tbl13 = doc.add_table(rows=1, cols=4)
tbl13.style = 'Table Grid'
table_hdr(tbl13, ['Label', 'CSV column', 'Unit', 'Notes'], [5.5, 4.0, 2.5, 3.5])
rows13 = [
    ('Agricultural Water Productivity (Awp)', 'awp_usd_per_m3', 'USD/m³',   'Core output'),
    ('Year-to-year change (cAwp)',            'cawp_pct',       '%',         'Undefined at 2018'),
    ('Trend since 2018 (tAwp)',               'tawp_pct',       '%',         '0 at 2018 by definition'),
    ('Blue ET volume (VETb)',                 'vetb_Mm3',       'Mm³/year',  'Stored as m³; displayed as Mm³'),
    ('Agricultural GVA',                      'gva_agriculture_usd', 'USD/yr', 'Includes forestry+fishing'),
    ('Irrigation-weighted area',              'irr_area_ha',    'ha',         'Source recorded in landuse_source column'),
    ('Annual blue ET depth (ETb)',            'ETb_annual_mm',  'mm/year',   'Area-weighted mean'),
    ('Annual AETI depth',                     'AETI_annual_mm', 'mm/year',   'Area-weighted mean'),
]
for i, r in enumerate(rows13):
    table_row(tbl13, r, i, [5.5, 4.0, 2.5, 3.5], alt=bool(i % 2))

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. KNOWN LIMITATIONS AND UNCERTAINTY
# ══════════════════════════════════════════════════════════════════════════════
h1('8.  Known Limitations and Uncertainty')

h2('8.1  GVA Includes Forestry and Fishing')
para(
    'World Bank NV.AGR.TOTL.CD aggregates crop farming, livestock, forestry, and '
    'fishing into a single national figure. For countries where forestry or fishing '
    'is a large share of agricultural GDP (e.g., Norway, Myanmar, Indonesia, Brazil), '
    'the Awp numerator is significantly overestimated.'
)
para(
    'Magnitude: countries where forestry+fishing > 20% of agricultural GVA '
    'may have Awp overestimated by 25% or more. These are flagged GVA_FORESTRY_FISHING.'
)
para(
    'Mitigation: The crop ratio Cr partially corrects for this by removing non-irrigated '
    'production from GVAa, but Cr itself is estimated from area proxies (Section 4.6), '
    'not from sector-level value data.'
)

h2('8.2  Crop Ratio (Cr) Uncertainty')
para(
    'Cr is the single largest source of Awp uncertainty for most countries. '
    'The area-proxy method (Cr ≈ 1 − irrigated/cultivated area) does not account for '
    'the higher value-per-hectare of irrigated crops relative to rainfed crops. '
    'This typically underestimates the irrigated value fraction, leading to Cr being '
    'slightly too high and Awp slightly underestimated.'
)
para(
    'For country-fixed and regional default Cr values (Levels 2–4), uncertainty is '
    'estimated at ±20–30% of the Cr value itself, propagating to ±10–25% uncertainty '
    'in Awp depending on the country\'s Cr level.'
)

h2('8.3  GMIE-100 Vintage and Accuracy')
para(
    'GMIE-100 represents irrigated area circa 2000. Since then, irrigation infrastructure '
    'has expanded significantly in parts of Sub-Saharan Africa, South Asia, and Central '
    'Asia, and contracted in some formerly Soviet states and parts of the Middle East '
    'due to conflict or economic change.'
)
para(
    'As a result, GMIE-100 may underestimate current irrigation in fast-growing regions '
    'and overestimate it in regions of abandonment. This affects both VETb (volume) '
    'and the GMIE-weighted area. A newer global irrigation map, when available, '
    'should be substituted.'
)

h2('8.4  WaPOR Coverage and Partial Years')
para(
    'WaPOR v3 L1 products start in 2018. Pre-2018 analysis is not possible with '
    'this pipeline without substituting an alternative ET product.'
)
para(
    'Near-real-time years may not yet have a complete 12-month record when the '
    'pipeline is run. Rows with n_months_aeti < 12 are flagged PARTIAL_YEAR, '
    'and their annual ETb sum (and therefore VETb and Awp) will be underestimated '
    'proportionally to the number of missing months.'
)

h2('8.5  Multi-Polygon Countries and Small Island States')
para(
    'LSIB_SIMPLE represents countries as individual polygons, including archipelagos '
    'as multiple separate features. The GEE reduceRegions step produces one result row '
    'per polygon, which Python then aggregates by ISO-3. For small island states where '
    'all polygons are small (< 1,000 ha GMIE area each), aggregation may still not '
    'reach the minimum threshold, resulting in LOW_GMIE_AREA or NO_GMIE_IRRIGATION flags.'
)

h2('8.6  Country ISO-3 Code Mismatches')
para(
    'Twelve countries have known ISO-3 alignment challenges across data sources '
    '(Kosovo XKX, Taiwan TWN, Western Sahara ESH, Palestine PSE, Kosovo XKX, '
    'South Sudan SSD, etc.). These are flagged REVIEW in the crosswalk and '
    'BOUNDARY_MISMATCH in the output if inconsistencies are detected. Results for '
    'these countries should be interpreted with caution.'
)

h2('8.7  Spatial Scale Mismatch Between GMIE and WaPOR')
para(
    'GMIE-100 is at ~1 km resolution while WaPOR v3 L1 products are at ~300 m. '
    'GMIE is bilinearly resampled to 300 m in the GEE pipeline. This means each '
    '300 m pixel receives an interpolated fraction from its surrounding ~1 km '
    'neighbourhood, which is physically appropriate for a continuously varying '
    'irrigation fraction but introduces smoothing at field boundaries.'
)

h2('8.8  What This System Does Not Capture')
bullet('Livestock water use (included in GVA but not in VETb)')
bullet('Groundwater depletion vs. surface water use (both count as blue ET)')
bullet('Irrigation efficiency losses before water reaches the field')
bullet('Subsistence farming with no formal irrigation but some field flooding')
bullet('Intra-annual variability (output is annual averages only)')

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 9. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1('9.  References')

refs = [
    ('Brouwer, C. & Heibloem, M. (1986)',
     'Irrigation Water Management: Irrigation Water Needs. Training Manual No. 3. '
     'Food and Agriculture Organization of the United Nations, Rome.'),

    ('FAO (2023)',
     'WaPOR v3 — Water Productivity through Open access of Remotely sensed derived data. '
     'FAO, Rome. https://www.fao.org/in-action/remote-sensing-for-water-productivity'),

    ('Siebert, S. et al. (2005)',
     'Development and validation of the global map of irrigation areas. '
     'Hydrology and Earth System Sciences, 9(5), 535–547.'),

    ('Thenkabail, P.S. et al. (2009)',
     'Global Irrigated Area Map (GIAM), derived from remote sensing, for the end of '
     'the last millennium. International Journal of Remote Sensing, 30(14), 3679–3733.'),

    ('World Bank (2024)',
     'World Development Indicators — NV.AGR.TOTL.CD: Agriculture, forestry and '
     'fishing, value added (current US$). data.worldbank.org'),

    ('FAO AQUASTAT (2024)',
     'AQUASTAT Main Database. Food and Agriculture Organization of the United Nations. '
     'www.fao.org/aquastat/en/databases'),

    ('Lorite, I.J. et al. (2004)',
     'Assessing deficit irrigation strategies at the level of an irrigation district. '
     'Agricultural Water Management, 91(1–3), 30–42.'),

    ('Molden, D. (ed.) (2007)',
     'Water for Food, Water for Life: A Comprehensive Assessment of Water Management '
     'in Agriculture. Earthscan, London; IWMI, Colombo.'),

    ('Zwart, S.J. & Bastiaanssen, W.G.M. (2004)',
     'Review of measured crop water productivity values for irrigated wheat, rice, '
     'cotton and maize. Agricultural Water Management, 69(2), 115–133.'),
]

for authors, text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after  = Pt(6)
    run_a = p.add_run(authors + '  ')
    run_a.bold = True
    run_a.font.size = Pt(10)
    run_t = p.add_run(text)
    run_t.font.size = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size:  {OUT.stat().st_size / 1024:.0f} KB')
